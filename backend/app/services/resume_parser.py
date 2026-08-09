import io
import re
import json
from typing import Dict, Any, List, Tuple
from app.services.ai_interview_engine import ai_interview_engine
from app.core.logging import logger
from app.schemas.resume import ParsedResume, ResumeExperience, ResumeProject, ResumeEducation

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

class ResumeParserService:
    
    @staticmethod
    async def parse_and_extract(resume_text: str, filename: str = "resume.pdf") -> Dict[str, Any]:
        """
        Legacy method for AI interview engine compatibility.
        """
        prompt = f"""
You are an expert technical recruiter analyzing a candidate's resume.

Resume Document Text:
\"\"\"{resume_text[:4000]}\"\"\"

Extract structured candidate insights and generate 5 tailored interview questions.

Respond ONLY in valid JSON format (no markdown code blocks):
{{
  "extracted_skills": ["Python", "React", "FastAPI", "MongoDB", "Docker", "AWS"],
  "extracted_projects": ["StudyForge AI Platform", "Real-Time Distributed Chat Service"],
  "extracted_experience": ["Software Development Engineer at Tech Corp (2 years)"],
  "generated_questions": [
    "I see you built a real-time chat service. How did you handle WebSocket connection failures and scaling?",
    "Can you explain your experience using FastAPI and MongoDB Atlas in production?",
    "Tell me about a technical trade-off you had to make when designing your StudyForge AI platform."
  ]
}}
"""
        raw = await ai_interview_engine._call_nvidia_nim([{"role": "user", "content": prompt}], temperature=0.3)
        try:
            cleaned = raw.strip()
            if cleaned.startswith("```json"):
                cleaned = cleaned[7:]
            if cleaned.startswith("```"):
                cleaned = cleaned[3:]
            if cleaned.endswith("```"):
                cleaned = cleaned[:-3]
            return json.loads(cleaned.strip())
        except Exception as e:
            logger.error(f"Failed to parse resume JSON: {e}")
            return {
                "extracted_skills": ["Software Engineering", "Full-Stack Development", "Problem Solving"],
                "extracted_projects": ["Full-Stack Web Application"],
                "extracted_experience": ["Software Engineer"],
                "generated_questions": [
                    "Can you walk me through your most complex software project?",
                    "How do you ensure high performance and low latency in your web applications?"
                ]
            }

    @staticmethod
    def detect_file_type(filename: str, content_type: str) -> str:
        filename = filename.lower()
        if filename.endswith(".pdf") or content_type == "application/pdf":
            return "pdf"
        if filename.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        return "unknown"

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> Tuple[str, bool, float, List[str]]:
        if not fitz:
            return "", False, 0.0, ["PyMuPDF not installed"]
        
        text = ""
        warnings = []
        ocr_used = False
        quality = 1.0
        
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text() + "\n\n"
            
            # Detect if scanned
            if len(text.strip()) < 100 and len(doc) > 0:
                warnings.append("Insufficient text extracted from PDF. Document might be scanned.")
                ocr_text = ResumeParserService.ocr_fallback(file_bytes)
                if ocr_text:
                    text = ocr_text
                    ocr_used = True
                    quality = 0.6
                else:
                    quality = 0.1
            
            doc.close()
            return text, ocr_used, quality, warnings
        except Exception as e:
            logger.error(f"PDF Parse Error: {e}")
            return "", False, 0.0, [f"PDF parse error: {str(e)}"]

    @staticmethod
    def parse_docx(file_bytes: bytes) -> Tuple[str, bool, float, List[str]]:
        if not Document:
            return "", False, 0.0, ["python-docx not installed"]
            
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_data:
                        paragraphs.append(" | ".join(row_data))
                        
            return "\n\n".join(paragraphs), False, 1.0, []
        except Exception as e:
            logger.error(f"DOCX Parse Error: {e}")
            return "", False, 0.0, [f"DOCX parse error: {str(e)}"]

    @staticmethod
    def ocr_fallback(file_bytes: bytes) -> str:
        if not fitz or not pytesseract or not Image:
            return ""
            
        try:
            text = ""
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n\n"
            doc.close()
            return text
        except Exception as e:
            logger.error(f"OCR Error: {e}")
            return ""

    @staticmethod
    def normalize_text(text: str) -> str:
        # Normalize whitespace and bullets
        text = re.sub(r'[\u2022\u2023\u25E6\u2043\u2219\u25CB\u25CF\u27A4\u27A2\u27A1\u2794\u2192]', '-', text)
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        return text.strip()

    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        # Basic section detection
        section_headers = {
            "summary": [r"summary", r"profile", r"objective", r"about me", r"professional summary"],
            "experience": [r"experience", r"work experience", r"employment history", r"professional experience", r"work history"],
            "education": [r"education", r"academic background"],
            "skills": [r"skills", r"technologies", r"technical skills", r"core competencies"],
            "projects": [r"projects", r"personal projects", r"academic projects", r"software projects"],
            "certifications": [r"certifications", r"licenses", r"courses"],
            "achievements": [r"achievements", r"awards", r"honors"],
            "languages": [r"languages"]
        }
        
        # A simple regex based parser
        lines = text.split('\n')
        sections = {"contact": "", "summary": "", "experience": "", "education": "", "skills": "", "projects": "", "certifications": "", "achievements": "", "languages": "", "uncategorized": ""}
        current_section = "contact" # First part is usually contact info
        
        for line in lines:
            line_clean = line.strip().lower()
            if not line_clean:
                sections[current_section] += "\n"
                continue
                
            matched_section = None
            if len(line_clean) < 40 and not any(char.islower() for char in line_clean) and len(line_clean.split()) < 5:
                # Potential uppercase header
                for key, patterns in section_headers.items():
                    for pattern in patterns:
                        if re.match(r"^" + pattern + r"s?$", line_clean):
                            matched_section = key
                            break
                    if matched_section:
                        break
            else:
                for key, patterns in section_headers.items():
                    for pattern in patterns:
                        if re.match(r"^" + pattern + r"s?$", line_clean):
                            matched_section = key
                            break
                    if matched_section:
                        break

            if matched_section:
                current_section = matched_section
            else:
                sections[current_section] += line + "\n"
                
        return sections

    @staticmethod
    def parse_resume_file(file_bytes: bytes, filename: str, content_type: str) -> ParsedResume:
        file_type = ResumeParserService.detect_file_type(filename, content_type)
        
        if file_type == "pdf":
            raw_text, ocr_used, quality, warnings = ResumeParserService.parse_pdf(file_bytes)
            parser = "PyMuPDF"
        elif file_type == "docx":
            raw_text, ocr_used, quality, warnings = ResumeParserService.parse_docx(file_bytes)
            parser = "python-docx"
        else:
            return ParsedResume(
                parser="unknown",
                parse_quality=0.0,
                warnings=["Unsupported file format. Please upload PDF or DOCX."]
            )
            
        if not raw_text.strip():
            return ParsedResume(
                parser=parser,
                ocr_used=ocr_used,
                parse_quality=0.0,
                warnings=["Document is empty or could not be parsed."] + warnings
            )
            
        normalized_text = ResumeParserService.normalize_text(raw_text)
        sections = ResumeParserService.extract_sections(normalized_text)
        
        # Populate basic structure (Note: Deeper NLP extraction to populate Experience/Projects properly would be phase 8+)
        # For now, just dumping section text into the parsed resume. We'll populate skills by line splitting for now.
        
        skills_list = []
        if sections["skills"]:
            skills_list = [s.strip() for s in re.split(r'[,\-|\n]', sections["skills"]) if s.strip() and len(s.strip()) > 1]
            
        return ParsedResume(
            contact={"raw": sections["contact"].strip()},
            summary=sections["summary"].strip(),
            skills=list(set(skills_list)),
            raw_text=normalized_text,
            parser=parser,
            ocr_used=ocr_used,
            parse_quality=quality,
            warnings=warnings
        )

resume_parser_service = ResumeParserService()
